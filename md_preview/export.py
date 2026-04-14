"""Export Markdown to PDF, DOCX, HTML, or browser preview."""

from __future__ import annotations

import re
import tempfile
import webbrowser
from pathlib import Path
from typing import Callable

from md_preview.render import extract_title, to_html_body
from md_preview.styles import GITHUB_CSS, HTML_WRAPPER


def to_html(text: str, output: Path) -> Path:
    """Render Markdown to a standalone HTML file with syntax highlighting."""
    body = to_html_body(text)
    title = extract_title(text)
    html = HTML_WRAPPER.format(title=title, css=GITHUB_CSS, body=body)
    output.write_text(html, encoding="utf-8")
    return output


def to_pdf(text: str, output: Path) -> Path:
    """Render Markdown to a paginated PDF."""
    from weasyprint import CSS, HTML

    body = to_html_body(text)
    title = extract_title(text)
    html_str = HTML_WRAPPER.format(title=title, css=GITHUB_CSS, body=body)
    HTML(string=html_str).write_pdf(str(output), stylesheets=[CSS(string=GITHUB_CSS)])
    return output


def to_docx(text: str, output: Path) -> Path:
    """Render Markdown to a Word document."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    _apply_default_style(doc)

    in_code_block = False
    code_buf: list[str] = []

    for line in text.split("\n"):
        if line.startswith("```"):
            if in_code_block:
                _add_code_block(doc, "\n".join(code_buf))
                code_buf.clear()
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_buf.append(line)
            continue

        s = line.strip()
        if not s:
            continue

        if s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("#### "):
            doc.add_heading(s[5:], level=4)
        elif s.startswith("---"):
            _add_hr(doc)
        elif s.startswith(("- ", "* ")):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", s):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", s), style="List Number")
        elif s.startswith("> "):
            _add_blockquote(doc, s[2:])
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, s)

    doc.save(str(output))
    return output


def to_md(text: str, output: Path) -> Path:
    """Copy raw Markdown to the output path."""
    output.write_text(text, encoding="utf-8")
    return output


def preview_in_browser(text: str) -> Path:
    """Render to a temp HTML file and open in the default browser."""
    tmp = Path(tempfile.gettempdir()) / "md-preview.html"
    to_html(text, tmp)
    webbrowser.open(tmp.as_uri())
    return tmp


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def _apply_default_style(doc) -> None:  # noqa: ANN001
    from docx.shared import Pt, RGBColor

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1F, 0x23, 0x28)


def _add_code_block(doc, code: str) -> None:  # noqa: ANN001
    from docx.shared import Pt

    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def _add_hr(doc) -> None:  # noqa: ANN001
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 50)
    run.font.color.rgb = RGBColor(0xD1, 0xD9, 0xE0)


def _add_blockquote(doc, text: str) -> None:  # noqa: ANN001
    from docx.shared import RGBColor

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x65, 0x6D, 0x76)


def _add_inline_runs(paragraph, text: str) -> None:  # noqa: ANN001
    """Parse **bold**, *italic*, and `code` spans into Word runs."""
    from docx.shared import Pt, RGBColor

    for part in re.split(r"(\*\*.*?\*\*|`.*?`|\*.*?\*)", text):
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part:
            paragraph.add_run(part)


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

FORMATS: dict[str, Callable[[str, Path], Path]] = {
    "html": to_html,
    "pdf": to_pdf,
    "docx": to_docx,
    "md": to_md,
}
