"""Export Markdown / rich turns to PDF, DOCX, HTML, or browser preview."""

from __future__ import annotations

import re
import tempfile
import webbrowser
from pathlib import Path
from typing import Callable

from preview.render import extract_title, to_html_body, to_rich_html_body
from preview.styles import BROWSER_CSS, COPY_BUTTON_JS, DOCUMENT_CSS, HTML_WRAPPER
from preview.session import RichTurn


def _rich_or_plain(
    turns_or_text: list[RichTurn] | str,
    *,
    rich: bool = True,
    session_id: str = "",
) -> tuple[str, str]:
    """Return (html_body, title) from either rich turns or plain text."""
    if rich and isinstance(turns_or_text, list):
        body = to_rich_html_body(turns_or_text, session_id=session_id)
        texts = " ".join(t.plain_text() for t in turns_or_text)
        title = extract_title(texts)
    else:
        text = turns_or_text if isinstance(turns_or_text, str) else "\n\n---\n\n".join(t.plain_text() for t in turns_or_text)
        body = to_html_body(text)
        title = extract_title(text)
    return body, title


# ---------------------------------------------------------------------------
# HTML / browser
# ---------------------------------------------------------------------------

def to_html(source: list[RichTurn] | str, output: Path, session_id: str = "") -> Path:
    """Render to a styled HTML file with tool events and copy buttons."""
    body, title = _rich_or_plain(source, session_id=session_id)
    page = HTML_WRAPPER.format(title=title, css=BROWSER_CSS, body=body, js=COPY_BUTTON_JS)
    output.write_text(page, encoding="utf-8")
    return output


def preview_in_browser(source: list[RichTurn] | str, session_id: str = "") -> Path:
    """Render to a temp HTML file and open in the default browser."""
    tmp = Path(tempfile.gettempdir()) / "preview.html"
    to_html(source, tmp, session_id=session_id)
    webbrowser.open(tmp.as_uri())
    return tmp


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def to_pdf(source: list[RichTurn] | str, output: Path) -> Path:
    """Render to a paginated PDF using DOCUMENT_CSS (no JS, no tool blocks)."""
    from weasyprint import CSS, HTML as WP_HTML

    # PDF gets plain text only — tool blocks don't translate well to print
    if isinstance(source, list):
        text = "\n\n---\n\n".join(t.plain_text() for t in source)
    else:
        text = source
    body = to_html_body(text)
    title = extract_title(text)
    html_str = HTML_WRAPPER.format(title=title, css=DOCUMENT_CSS, body=body, js="")
    WP_HTML(string=html_str, base_url="/").write_pdf(str(output), stylesheets=[CSS(string=DOCUMENT_CSS)])
    return output


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def to_docx(source: list[RichTurn] | str, output: Path) -> Path:
    """Render Markdown to a Word document."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    text = "\n\n---\n\n".join(t.plain_text() for t in source) if isinstance(source, list) else source
    doc = Document()
    _apply_default_style(doc)

    in_code_block = False
    code_buf: list[str] = []

    for line in text.split("\n"):
        if line.startswith("```"):
            if in_code_block:
                _add_code_block(doc, "\n".join(code_buf))
                code_buf.clear()
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_buf.append(line)
            continue

        s = line.strip()
        if not s:
            continue

        if s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("#### "):
            doc.add_heading(s[5:], level=4)
        elif s.startswith("---"):
            _add_hr(doc)
        elif s.startswith(("- ", "* ")):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", s):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", s), style="List Number")
        elif s.startswith("> "):
            _add_blockquote(doc, s[2:])
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, s)

    doc.save(str(output))
    return output


def to_md(source: list[RichTurn] | str, output: Path) -> Path:
    text = "\n\n---\n\n".join(t.plain_text() for t in source) if isinstance(source, list) else source
    output.write_text(text, encoding="utf-8")
    return output


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def _apply_default_style(doc) -> None:
    from docx.shared import Pt, RGBColor
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1F, 0x23, 0x28)


def _add_code_block(doc, code: str) -> None:
    from docx.shared import Pt
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def _add_hr(doc) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 50)
    run.font.color.rgb = RGBColor(0xD1, 0xD9, 0xE0)


def _add_blockquote(doc, text: str) -> None:
    from docx.shared import RGBColor
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x65, 0x6D, 0x76)


def _add_inline_runs(paragraph, text: str) -> None:
    from docx.shared import Pt, RGBColor
    for part in re.split(r"(\*\*.*?\*\*|`.*?`|\*.*?\*)", text):
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part:
            paragraph.add_run(part)


# ---------------------------------------------------------------------------
# Format registry
# ---------------------------------------------------------------------------

FORMATS: dict[str, Callable[[list[RichTurn] | str, Path], Path]] = {
    "html": to_html,
    "pdf": to_pdf,
    "docx": to_docx,
    "md": to_md,
}
